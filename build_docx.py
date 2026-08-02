"""Erzeugt aus den Markdown-Referenzdokumenten in Basisinfos/ die .docx-Lesekopien.

Hintergrund (2026-08-02): die .docx-Varianten wurden bisher von Hand nachgezogen
und sind dabei mehrfach hinter der .md zurueckgeblieben - beim Regelwerksmanual-
Split vom 02.08. am deutlichsten (die .docx enthielt noch die ungeteilte Fassung).
Die .md-Dateien sind und bleiben die Quelle der Wahrheit (git-versioniert); die
.docx sind reine Lesekopien zum Durchblaettern/Drucken und werden hier jederzeit
reproduzierbar neu erzeugt.

Aufruf:
    python build_docx.py            # alle konfigurierten Dokumente
    python build_docx.py --pruefen  # nur melden, welche .docx veraltet sind
    python build_docx.py Regelwerksmanual.md

Bewusst KEIN Scheduler-Job: die .docx werden selten gebraucht und ein
automatischer Lauf wuerde bei jedem Doku-Commit binaere Dateien mitaendern.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASIS = Path(__file__).parent / 'Basisinfos'

# Welche .md eine .docx-Lesekopie bekommen. Bewusst nicht "alle .md": z.B.
# MEMORY_SYNC_STRATEGY.md und Regler_Signal_Pipeline_Abhaengigkeiten.md sind
# reine Arbeitsdokumente, die nur im Editor gelesen werden.
DOKUMENTE = [
    'Regelwerksmanual.md',
    'Regelwerk_Entscheidungslog.md',
    'Test_und_Verifikationsmethodik.md',
    'Fakten_Entscheidungsmappe.md',
    'Kategorie_Basisinformationen_Release2.md',
    'Spezifikation.md',
]

CODE_FARBE = RGBColor(0x88, 0x22, 0x22)
CODE_FONT = 'Consolas'


# --------------------------------------------------------------- Inline-Format
# Markdown-Inline-Auszeichnung in Runs uebersetzen. Bewusst simpel gehalten -
# verschachtelte Auszeichnung (**fett mit `code` drin**) kommt in diesen
# Dokumenten nicht vor und wuerde einen echten Parser erfordern.
_INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)')


def _inline_runs(paragraph, text: str) -> None:
    for teil in _INLINE.split(text):
        if not teil:
            continue
        if teil.startswith('**') and teil.endswith('**'):
            run = paragraph.add_run(teil[2:-2])
            run.bold = True
        elif teil.startswith('`') and teil.endswith('`'):
            run = paragraph.add_run(teil[1:-1])
            run.font.name = CODE_FONT
            run.font.color.rgb = CODE_FARBE
            run.font.size = Pt(9)
        elif teil.startswith('*') and teil.endswith('*') and len(teil) > 2:
            run = paragraph.add_run(teil[1:-1])
            run.italic = True
        else:
            paragraph.add_run(teil)


def _ist_tabellen_trenner(zeile: str) -> bool:
    """Erkennt die |---|---|-Zeile unter einem Tabellenkopf."""
    kern = zeile.strip().strip('|')
    return bool(kern) and all(
        set(z.strip()) <= set('-: ') and '-' in z for z in kern.split('|')
    )


def _tabellen_zellen(zeile: str) -> list[str]:
    return [z.strip() for z in zeile.strip().strip('|').split('|')]


# ------------------------------------------------------------------ Konverter
def md_nach_docx(md_pfad: Path, docx_pfad: Path) -> dict:
    zeilen = md_pfad.read_text(encoding='utf-8').split('\n')
    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)

    stat = {'ueberschriften': 0, 'tabellen': 0, 'codebloecke': 0, 'absaetze': 0}
    i = 0
    while i < len(zeilen):
        zeile = zeilen[i]
        strip = zeile.strip()

        # Codeblock
        if strip.startswith('```'):
            i += 1
            code = []
            while i < len(zeilen) and not zeilen[i].strip().startswith('```'):
                code.append(zeilen[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('\n'.join(code))
            run.font.name = CODE_FONT
            run.font.size = Pt(8.5)
            stat['codebloecke'] += 1
            continue

        # Tabelle: Kopfzeile + |---|-Trenner
        if (strip.startswith('|') and i + 1 < len(zeilen)
                and _ist_tabellen_trenner(zeilen[i + 1])):
            kopf = _tabellen_zellen(zeile)
            i += 2
            reihen = []
            while i < len(zeilen) and zeilen[i].strip().startswith('|'):
                reihen.append(_tabellen_zellen(zeilen[i]))
                i += 1
            tabelle = doc.add_table(rows=1, cols=len(kopf))
            tabelle.style = 'Light Grid Accent 1'
            tabelle.alignment = WD_TABLE_ALIGNMENT.LEFT
            for spalte, titel in enumerate(kopf):
                zelle = tabelle.rows[0].cells[spalte]
                zelle.text = ''
                _inline_runs(zelle.paragraphs[0], titel)
                for run in zelle.paragraphs[0].runs:
                    run.bold = True
            for reihe in reihen:
                zellen = tabelle.add_row().cells
                for spalte, wert in enumerate(reihe[:len(kopf)]):
                    zellen[spalte].text = ''
                    _inline_runs(zellen[spalte].paragraphs[0], wert)
            doc.add_paragraph()
            stat['tabellen'] += 1
            continue

        # Ueberschrift
        m = re.match(r'^(#{1,6})\s+(.*)$', strip)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
            stat['ueberschriften'] += 1
            i += 1
            continue

        # Horizontale Linie
        if strip in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.add_run('─' * 60).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            i += 1
            continue

        # Blockzitat
        if strip.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _inline_runs(p, strip[2:])
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Aufzaehlung / nummerierte Liste (Einrueckung als Ebene beibehalten)
        m = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)$', zeile)
        if m:
            ebene = min(len(m.group(1)) // 2, 2)
            stil = 'List Number' if m.group(2).endswith('.') else 'List Bullet'
            if ebene:
                stil = f'{stil} {ebene + 1}'
            try:
                p = doc.add_paragraph(style=stil)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            # Umgebrochene Listenpunkte zusammenfuehren: eine eingerueckte
            # Folgezeile, die selbst kein neuer Listenpunkt ist, gehoert noch
            # zum selben Punkt. Ohne das zerfaellt in listenlastigen Dokumenten
            # jeder mehrzeilige Punkt in mehrere Absaetze.
            text = m.group(3)
            i += 1
            while (i < len(zeilen) and zeilen[i].strip()
                   and zeilen[i][:1] in (' ', '\t')
                   and not re.match(r'^\s*([-*+]|\d+\.)\s', zeilen[i])
                   and not zeilen[i].strip().startswith('```')):
                text += ' ' + zeilen[i].strip()
                i += 1
            _inline_runs(p, text)
            continue

        # Leerzeile
        if not strip:
            i += 1
            continue

        # Normaler Absatz - Folgezeilen bis zur naechsten Leerzeile anhaengen
        absatz = [zeile.rstrip()]
        i += 1
        while i < len(zeilen) and zeilen[i].strip() and not re.match(
                r'^(#{1,6}\s|```|\s*([-*+]|\d+\.)\s|>\s|\|)', zeilen[i]):
            absatz.append(zeilen[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _inline_runs(p, ' '.join(absatz))
        stat['absaetze'] += 1

    doc.save(docx_pfad)
    return stat


def main() -> int:
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    nur_pruefen = '--pruefen' in sys.argv
    ziele = args or DOKUMENTE

    veraltet = 0
    for name in ziele:
        md = BASIS / name
        docx = md.with_suffix('.docx')
        if not md.exists():
            print(f'FEHLT   {name}')
            continue

        if nur_pruefen:
            if not docx.exists():
                print(f'FEHLT   {docx.name} (noch nie erzeugt)')
                veraltet += 1
            elif docx.stat().st_mtime < md.stat().st_mtime:
                print(f'ALT     {docx.name}')
                veraltet += 1
            else:
                print(f'AKTUELL {docx.name}')
            continue

        stat = md_nach_docx(md, docx)
        print(f'OK      {docx.name:<45} '
              f"{stat['ueberschriften']:>4} Ueberschr. "
              f"{stat['absaetze']:>5} Abs. "
              f"{stat['tabellen']:>3} Tab. "
              f"{stat['codebloecke']:>3} Code")

    if nur_pruefen and veraltet:
        print(f'\n{veraltet} .docx veraltet - "python build_docx.py" erzeugt sie neu.')
        return 1
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
